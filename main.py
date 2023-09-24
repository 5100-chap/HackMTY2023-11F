# Importacion de librerias para pagina principal
import streamlit as st
import pandas as pd
import fitz  # Importar PyMuPDF (Fitz)
from docx import Document
from back import *


# Declaración de clase para el uso de la app web
myclass = TextAssistant()

band = 0
st.markdown(
    "<h1 style='text-align: center;'>TinyWrite.gpt</h1>", unsafe_allow_html=True
)

# Agregar la barra de navegación a la barra lateral
st.sidebar.title("Menú de herramientas")


# Agregar enlaces a diferentes páginas o secciones
pagina_inicio = st.sidebar.button("Correción de textos")
pagina_seccion1 = st.sidebar.button("Sugerencias de mejora de textos")
pagina_seccion2 = st.sidebar.button("Traductor de documentos")

if "current_section" not in st.session_state:
    st.session_state.current_section = "inicio"


# Función para mostrar el contenido de la página de inicio
def mostrar_inicio():
    st.empty()

    st.markdown(
        "<h3 style='text-align: left;'>Corrección de textos</h3>",
        unsafe_allow_html=True,
    )

    user_input = st.text_input("Introduzca su texto aquí:")

    # Crear un checkbox en Streamlit
    show_content = st.checkbox("¿Desea adjuntar un archivo?")

    st.write("\n".join(user_input))

    # Mostrar contenido si el checkbox está marcado
    if show_content:
        # Crear un campo de entrada de archivos en Streamlit

        uploaded_file = st.file_uploader(
            "Cargar un archivo:", type=["txt", "pdf", "docx"]
        )

        if uploaded_file is not None:
            st.markdown(
                "<h4 style='text-align: left; text-decoration: underline'> Texto original: </h4>",
                unsafe_allow_html=True,
            )

            # Inicializa una variable para almacenar el contenido del archivo
            contenido = []

            # Aquí puedes realizar operaciones con el archivo cargado, como leer su contenido o procesarlo.

            file_extension = uploaded_file.name.split(".")[-1].lower()

            if file_extension == "pdf":
                # Leer el contenido del archivo PDF
                pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")

                # Crear una lista para almacenar el contenido de cada página del PDF
                pdf_text_array = []

                # Iterar a través de las páginas del PDF y agregar el texto de cada página a la lista
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    pdf_text_array.append(page.get_text())

                # Mostrar el contenido del archivo en Streamlit como una lista
                st.write(pdf_text_array)
                print("------------------------------------------------")
                print(pdf_text_array)
                print("------------------------------------------------")
                
            elif file_extension == "docx":
                pdf_text_array = None
                # Leer el contenido del archivo docx
                doc = Document(uploaded_file)
                
                # Concatenar todos los párrafos en un solo string utilizando '\n' como delimitador
                full_text = '\n'.join([parrafo.text for parrafo in doc.paragraphs if parrafo.text.strip() != ''])
                
                # Crear una lista que contiene el texto completo (similar a pdf_text_array)
                doc_text_array = [full_text]

                # Mostrar el contenido del archivo en Streamlit como una lista
                st.write(doc_text_array)
                print("------------------------------------------------")
                print(doc_text_array)
                print("------------------------------------------------")


            else:
                # Lee el contenido del archivo TXT línea por línea
                for linea_bytes in uploaded_file:
                    # Decodifica la línea y elimina caracteres de puntuación
                    linea = linea_bytes.decode("utf-8").strip("\n")
                    contenido.append(linea)
 
            # Personalizar el st.write con CSS
            st.markdown(
                """
                <style>
                /* Estilos personalizados para st.write */
                .custom-write {
                    font-size: 24px; /* Tamaño de fuente */
                    color: #000000; /* Color de texto */
                    background-color: #f0f0f0; /* Color de fondo */
                    padding: 10px; /* Espaciado interno */
                    border: 2px solid #333; /* Borde */
                    border-radius: 10px; /* Bordes redondeados */
                    box-shadow: 3px 3px 5px 0px rgba(0,0,0,0.3); /* Sombra */
                }
                </style>
                """,
                unsafe_allow_html=True,
            )
            
            if pdf_text_array is not None:
                contenido = pdf_text_array
            else:
                contenido = doc_text_array
                st.write("\n".join(contenido))
            print(contenido)

            

            st.markdown(
                "<h4 style='text-align: left; text-decoration: underline'> Texto corregido: </h4>",
                unsafe_allow_html=True,
            )
            corrected = myclass.correct_text(contenido)
            print(corrected)
            st.write(corrected)


# Función para mostrar el contenido de la Sección 1
def mostrar_Sugerencia():
    st.empty()

    st.markdown(
        "<h3 style='text-align: left;'> Sugerencias de mejora de redacción de textos</h3>",
        unsafe_allow_html=True,
    )

    user_input = st.text_input("Introduzca su texto aquí:")

    # Crear un checkbox en Streamlit
    show_content = st.checkbox("¿Desea adjuntar un archivo?")

    st.write("\n".join(user_input))

    # Mostrar contenido si el checkbox está marcado
    if show_content:
        # Crear un campo de entrada de archivos en Streamlit

        uploaded_file = st.file_uploader(
            "Cargar un archivo:", type=["txt", "pdf", "docx"]
        )

        if uploaded_file is not None:
            st.markdown(
                "<h4 style='text-align: left; text-decoration: underline'> Texto original: </h4>",
                unsafe_allow_html=True,
            )

            # Inicializa una variable para almacenar el contenido del archivo
            contenido = []

            # Aquí puedes realizar operaciones con el archivo cargado, como leer su contenido o procesarlo.

            file_extension = uploaded_file.name.split(".")[-1].lower()

            if file_extension == "pdf":
                # Leer el contenido del archivo PDF
                pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")

                # Crear una lista para almacenar el contenido de cada página del PDF
                pdf_text_array = []

                # Iterar a través de las páginas del PDF y agregar el texto de cada página a la lista
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    pdf_text_array.append(page.get_text())

                # Mostrar el contenido del archivo en Streamlit como una lista
                st.write(pdf_text_array)
                print("------------------------------------------------")
                print(pdf_text_array)
                print("------------------------------------------------")
                
            elif file_extension == "docx":
                pdf_text_array = None
                # Leer el contenido del archivo docx
                doc = Document(uploaded_file)
                
                # Concatenar todos los párrafos en un solo string utilizando '\n' como delimitador
                full_text = '\n'.join([parrafo.text for parrafo in doc.paragraphs if parrafo.text.strip() != ''])
                
                # Crear una lista que contiene el texto completo (similar a pdf_text_array)
                doc_text_array = [full_text]

                # Mostrar el contenido del archivo en Streamlit como una lista
                st.write(doc_text_array)
                print("------------------------------------------------")
                print(doc_text_array)
                print("------------------------------------------------")


            else:
                # Lee el contenido del archivo TXT línea por línea
                for linea_bytes in uploaded_file:
                    # Decodifica la línea y elimina caracteres de puntuación
                    linea = linea_bytes.decode("utf-8").strip("\n")
                    contenido.append(linea)
 
            # Personalizar el st.write con CSS
            st.markdown(
                """
                <style>
                /* Estilos personalizados para st.write */
                .custom-write {
                    font-size: 24px; /* Tamaño de fuente */
                    color: #000000; /* Color de texto */
                    background-color: #f0f0f0; /* Color de fondo */
                    padding: 10px; /* Espaciado interno */
                    border: 2px solid #333; /* Borde */
                    border-radius: 10px; /* Bordes redondeados */
                    box-shadow: 3px 3px 5px 0px rgba(0,0,0,0.3); /* Sombra */
                }
                </style>
                """,
                unsafe_allow_html=True,
            )
            
            if pdf_text_array is not None:
                contenido = pdf_text_array
            else:
                contenido = doc_text_array
                st.write("\n".join(contenido))
            print(contenido)

            

            st.markdown(
                "<h4 style='text-align: left; text-decoration: underline'> Sugerencias: </h4>",
                unsafe_allow_html=True,
            )
            corrected = myclass.suggestText(contenido)
            print(corrected)
            st.write(corrected)


                

            

    


# Función para mostrar el contenido de la Sección 2
def mostrar_Traducción():
    st.empty()
    st.markdown(
        "<h3 style='text-align: left;'> Traducción de textos </h3>",
        unsafe_allow_html=True,
    )
    user_input = st.text_input("Coloque su idioma de selección:")

    # Inicializa una variable para almacenar el contenido del archivo o del input del usuario
    contenido = [user_input]

    # Crear un checkbox en Streamlit
    show_content = st.checkbox("¿Desea adjuntar un archivo?")

    if show_content and user_input!="":
        uploaded_file = st.file_uploader(
            "Cargar un archivo:", type=["txt", "pdf", "docx"]
        )

        if uploaded_file is not None:
            st.markdown(
                "<h4 style='text-align: left; text-decoration: underline'> Texto original: </h4>",
                unsafe_allow_html=True,
            )

            file_extension = uploaded_file.name.split(".")[-1].lower()

            if file_extension == "pdf":
                pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                contenido = [page.get_text() for page_num in range(len(pdf_document)) for page in [pdf_document[page_num]]]
            
            elif file_extension == "docx":
                doc = Document(uploaded_file)
                full_text = '\n'.join([parrafo.text for parrafo in doc.paragraphs if parrafo.text.strip() != ''])
                contenido = [full_text]

            elif file_extension == "txt":
                contenido = [linea_bytes.decode("utf-8").strip("\n") for linea_bytes in uploaded_file]

            st.write(contenido[0])
            # Traducir
            st.markdown(
                "<h4 style='text-align: left; text-decoration: underline'> Texto traducido: </h4>",
                unsafe_allow_html=True,
            )
            print(user_input)
            traduccion = myclass.traductor(contenido, user_input)
            st.write(traduccion)




if st.session_state.current_section == "inicio":
    mostrar_inicio()
elif st.session_state.current_section == "seccion1":
    mostrar_Sugerencia()
elif st.session_state.current_section == "seccion2":
    mostrar_Traducción()


if pagina_inicio:
    st.session_state.current_section = "inicio"
elif pagina_seccion1:
    st.session_state.current_section = "seccion1"
elif pagina_seccion2:
    st.session_state.current_section = "seccion2"

else:
    st.empty()